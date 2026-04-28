"""Parse document files into plain text for indexing."""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".html", ".htm", ".docx"}


@dataclass
class ParsedDocument:
    """A parsed document with extracted text and metadata."""

    path: Path
    text: str
    source_type: str  # "markdown" | "pdf" | "html" | "docx" | "txt"
    title: str  # from first heading or filename


def parse_document(path: Path) -> ParsedDocument:
    """Parse a document file into plain text. Raises ValueError for unsupported types."""
    suffix = path.suffix.lower()
    if suffix in (".md", ".txt"):
        text = path.read_text(encoding="utf-8")
        title = _extract_title(text, path)
        source_type = "markdown" if suffix == ".md" else "txt"
    elif suffix == ".pdf":
        text, title = _parse_pdf(path)
        source_type = "pdf"
    elif suffix in (".html", ".htm"):
        text, title = _parse_html(path)
        source_type = "html"
    elif suffix == ".docx":
        text, title = _parse_docx(path)
        source_type = "docx"
    else:
        raise ValueError(f"Unsupported file type: {suffix}")
    return ParsedDocument(path=path, text=text, source_type=source_type, title=title)


def _extract_title(text: str, path: Path) -> str:
    """Extract title from first markdown heading, or use filename."""
    for line in text.split("\n")[:10]:
        line = line.strip()
        if line.startswith("# "):
            return line[2:].strip()
    return path.stem


def _parse_pdf(path: Path) -> tuple[str, str]:
    """Extract text from a PDF using PyMuPDF."""
    import fitz  # pymupdf

    doc = fitz.open(str(path))
    pages = [page.get_text() for page in doc]
    doc.close()
    text = "\n\n".join(pages)
    title = _extract_title(text, path)
    return text, title


def _parse_html(path: Path) -> tuple[str, str]:
    """Extract text from an HTML file using BeautifulSoup."""
    from bs4 import BeautifulSoup

    html = path.read_text(encoding="utf-8")
    soup = BeautifulSoup(html, "html.parser")
    title_tag = soup.find("title")
    title = title_tag.get_text().strip() if title_tag else path.stem
    text = soup.get_text(separator="\n")
    return text, title


def _parse_docx(path: Path) -> tuple[str, str]:
    """Extract text from a .docx file using python-docx."""
    from docx import Document

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs]
    text = "\n\n".join(paragraphs)
    title = _extract_title(text, path)
    return text, title
